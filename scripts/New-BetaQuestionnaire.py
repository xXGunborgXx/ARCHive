from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "beta" / "ARCHive_Beta_Test_Questionnaire.docx"
LOGO = ROOT / "installer" / "ARCHive-wizard-small.png"

GRAPHITE = "1B2028"
GRAPHITE_2 = "343B47"
AMBER = "C47A32"
AMBER_LIGHT = "F7E6D3"
MUTED = "6B7280"
LIGHT = "F3F4F6"
WHITE = "FFFFFF"
INK = "20242B"
BORDER = "CDD2D9"

# compact_reference_guide with a named form-space override:
# Letter portrait, 0.75-inch margins, Calibri 11 pt, 1.25 line spacing,
# 6 pt body after, explicit fixed-width form tables and cell margins.
PAGE_WIDTH_DXA = 12240
MARGIN_DXA = 1080
CONTENT_DXA = PAGE_WIDTH_DXA - (MARGIN_DXA * 2)
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 90, "bottom": 90, "start": 120, "end": 120}


def set_run_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text, bold=False, color=INK, size=10.5, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(
        r,
        size={1: 16, 2: 13, 3: 12}[level],
        bold=True,
        color=GRAPHITE if level == 1 else AMBER,
    )
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        set_run_font(p.add_run(text))
    return p


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    shade(cell, AMBER_LIGHT)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(title), size=10.5, bold=True, color=GRAPHITE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_run_font(p2.add_run(text), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_form_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, prompt in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, color=GRAPHITE, size=10)
        shade(cells[0], LIGHT)
        set_cell_text(cells[1], prompt, color=MUTED, size=10, italic=True)
    set_table_geometry(table, [2500, CONTENT_DXA - 2500])
    return table


def add_workflow_table(doc):
    headers = ["Test", "Not tested", "Passed", "Problem", "Notes"]
    tests = [
        "Installer and disclosure",
        "First launch and expiry notice",
        "Copy one file",
        "Copy one folder",
        "Mixed files and folders",
        "Large-file copy",
        "Many-small-files copy",
        "Pause and Resume",
        "Cancel and cleanup",
        "Create 7z archive",
        "Create ZIP archive",
        "Extract 7z archive",
        "Extract ZIP archive",
        "Open Destination",
        "Diagnostic Log",
    ]
    widths = [3150, 1050, 900, 950, CONTENT_DXA - 6050]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell, GRAPHITE)
        set_cell_text(cell, header, bold=True, color=WHITE, size=9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for name in tests:
        cells = table.add_row().cells
        set_cell_text(cells[0], name, size=9.5)
        for index in (1, 2, 3):
            set_cell_text(cells[index], "[ ]", size=10)
            cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_text(cells[4], "", size=9.5)
    set_table_geometry(table, widths)
    return table


def add_rating_table(doc):
    headers = ["Area", "1", "2", "3", "4", "5", "Comment"]
    items = [
        "Installer clarity",
        "UI appearance",
        "Text readability",
        "Selecting sources",
        "Progress information",
        "Speed display",
        "Pause/Cancel confidence",
        "Result messages",
        "Overall ease of use",
    ]
    widths = [2850, 500, 500, 500, 500, 500, CONTENT_DXA - 5350]
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        shade(cell, GRAPHITE)
        set_cell_text(cell, header, bold=True, color=WHITE, size=9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for item in items:
        cells = table.add_row().cells
        set_cell_text(cells[0], item, size=9.5)
        for index in range(1, 6):
            set_cell_text(cells[index], "[ ]", size=9.5)
            cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_text(cells[6], "", size=9.5)
    set_table_geometry(table, widths)
    return table


def add_response_box(doc, prompt, lines=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(prompt), size=10.5, bold=True, color=GRAPHITE)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    shade(cell, "FBFBFC")
    cell.text = ""
    for idx in range(lines):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        if idx == 0:
            set_run_font(
                paragraph.add_run("Type your response here."),
                size=9.5,
                color=MUTED,
                italic=True,
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, before, after in ((1, 18, 10), (2, 14, 7), (3, 10, 5)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt({1: 16, 2: 13, 3: 12}[level])
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(
            GRAPHITE if level == 1 else AMBER
        )
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(7))
    set_table_geometry(table, [5000, CONTENT_DXA - 5000])
    set_cell_text(
        table.cell(0, 0),
        "ARCHive Beta Feedback",
        bold=True,
        color=GRAPHITE,
        size=9,
    )
    set_cell_text(
        table.cell(0, 1),
        "Seven-Day Beta • 1.0.0-beta1",
        color=MUTED,
        size=9,
    )
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in table.rows[0].cells:
        cell._tc.get_or_add_tcPr().remove_all("w:tcBorders")

    add_page_number(section.footer.paragraphs[0])


def build():
    doc = Document()
    configure_document(doc)

    # Page 1: customer-pack title block and tester profile.
    logo = doc.add_picture(str(LOGO), width=Inches(0.62))
    logo._inline.docPr.set("title", "ARCHive application logo")
    logo._inline.docPr.set(
        "descr",
        "ARCHive shield and storage application logo",
    )
    logo_paragraph = doc.paragraphs[-1]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_paragraph.paragraph_format.space_after = Pt(4)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(1)
    set_run_font(
        kicker.add_run("BETA TESTER PACK"),
        size=10,
        bold=True,
        color=AMBER,
    )
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(
        title.add_run("ARCHive Beta Test Questionnaire"),
        size=26,
        bold=True,
        color=GRAPHITE,
    )
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    set_run_font(
        subtitle.add_run(
            "Copy, archive, and extraction feedback • Seven-day beta"
        ),
        size=12.5,
        color=MUTED,
    )

    meta = doc.add_table(rows=2, cols=2)
    meta.style = "Table Grid"
    data = [
        ("Version", "1.0.0-beta1"),
        ("Return to", "GunborgServers@gmail.com"),
        ("Trial", "Seven days from first launch"),
        ("Format", "Edit in Word or compatible software"),
    ]
    for row_idx in range(2):
        for col_idx in range(2):
            label, value = data[(row_idx * 2) + col_idx]
            cell = meta.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            set_run_font(p.add_run(label + ": "), size=9.5, bold=True)
            set_run_font(p.add_run(value), size=9.5, color=MUTED)
    set_table_geometry(meta, [CONTENT_DXA // 2, CONTENT_DXA // 2])

    add_note(
        doc,
        "How to return this form",
        "Type directly into the fields, save a copy, and email it to "
        "GunborgServers@gmail.com. Screenshots, videos, and diagnostic JSON "
        "files are optional. Review diagnostic files first because they may "
        "contain complete file and folder paths.",
    )

    add_heading(doc, "Tester and Environment", 1)
    add_form_table(
        doc,
        [
            ("Name or tester alias", "Optional"),
            ("Reply email", "Optional"),
            ("Testing dates", "Start and end dates"),
            ("Windows edition/version", "Example: Windows 11 Pro 24H2"),
            ("Machine type", "[ ] Physical PC    [ ] VM"),
            ("Processor and RAM", "Example: Ryzen 7, 32 GB"),
            ("Display scaling", "Example: 100%, 125%, or 150%"),
            ("Source storage", "HDD / SSD / USB / network / VM shared folder"),
            ("Destination storage", "HDD / SSD / USB / network / VM shared folder"),
        ],
    )

    doc.add_page_break()

    # Page 2: workflow matrix.
    add_heading(doc, "Workflow Results", 1)
    add_body(
        doc,
        "Mark one result for every test you attempted. Use the Notes column "
        "for file size, storage type, unexpected delays, or error wording.",
    )
    add_workflow_table(doc)
    add_note(
        doc,
        "Integrity check",
        "Did every completed output open correctly or match the source? "
        "[ ] Yes  [ ] No  [ ] Not checked. If No, describe the exact item in "
        "the issue section and stop using that output as a backup.",
    )
    add_form_table(
        doc,
        [
            ("Largest single file", "Size and file type"),
            ("Largest full operation", "Total size and item count"),
            ("Typical copy speed", "Observed range and storage path"),
            ("Typical archive speed", "Observed range and compression setting"),
        ],
    )

    doc.add_page_break()

    # Page 3: ratings and behavior.
    add_heading(doc, "Usability Ratings", 1)
    add_body(
        doc,
        "Mark one score per row: 1 = very poor, 3 = acceptable, "
        "5 = excellent.",
    )
    add_rating_table(doc)
    add_heading(doc, "Progress and Waiting Behavior", 2)
    add_form_table(
        doc,
        [
            (
                "Progress accuracy",
                "Did percentage movement match visible file progress?",
            ),
            (
                "Speed display",
                "Was the number understandable and reasonably stable?",
            ),
            (
                "Waiting for storage",
                "Did this message prevent you from thinking ARCHive froze?",
            ),
            (
                "Pause/Resume",
                "Did the state and integrity behavior feel clear?",
            ),
            (
                "Cancellation",
                "Was cleanup and the final message understandable?",
            ),
        ],
    )
    add_response_box(
        doc,
        "What part of ARCHive gave you the most confidence?",
        lines=2,
    )
    add_response_box(
        doc,
        "What part felt confusing, slow, or unnecessary?",
        lines=2,
    )

    doc.add_page_break()

    # Page 4: issue report and final recommendation.
    add_heading(doc, "Most Important Issue", 1)
    add_form_table(
        doc,
        [
            ("Issue title", "Short description"),
            ("Action", "Copy / Create Archive / Extract Archive / Installer"),
            ("Severity", "[ ] Cosmetic  [ ] Minor  [ ] Major  [ ] Data integrity"),
            ("Frequency", "[ ] Once  [ ] Sometimes  [ ] Every time"),
            ("File workload", "Item count, total size, and file types"),
            ("Source → destination", "Storage types; avoid private path details"),
        ],
    )
    add_response_box(
        doc,
        "Steps to reproduce the issue",
        lines=4,
    )
    add_response_box(
        doc,
        "What happened, and what did you expect instead?",
        lines=4,
    )
    add_response_box(
        doc,
        "Exact message shown by ARCHive or Windows",
        lines=2,
    )

    doc.add_page_break()

    # Page 5: final feedback.
    add_heading(doc, "Final Feedback", 1)
    add_response_box(
        doc,
        "Which feature was most useful, and why?",
        lines=3,
    )
    add_response_box(
        doc,
        "What should be improved before a public release?",
        lines=4,
    )
    add_response_box(
        doc,
        "Was anything missing that an ordinary user would expect?",
        lines=3,
    )
    add_form_table(
        doc,
        [
            (
                "Would you use ARCHive again?",
                "[ ] Yes  [ ] Maybe  [ ] No",
            ),
            (
                "Ready for wider testing?",
                "[ ] Yes  [ ] After fixes  [ ] No",
            ),
            (
                "May we contact you?",
                "[ ] Yes  [ ] No",
            ),
            (
                "Attachments included",
                "[ ] Screenshot  [ ] Video  [ ] Diagnostic JSON  [ ] None",
            ),
        ],
    )
    add_note(
        doc,
        "Submit your feedback",
        "Save this completed document with your name or tester alias in the "
        "filename and email it to GunborgServers@gmail.com. Thank you for "
        "helping make ARCHive safer and easier to use.",
    )

    core_props = doc.core_properties
    core_props.title = "ARCHive Beta Test Questionnaire"
    core_props.subject = "Seven-day beta feedback form"
    core_props.author = "ARCHive Project"
    core_props.keywords = "ARCHive, beta, testing, questionnaire"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
